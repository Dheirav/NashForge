"""Build the NashForge report in IEEE conference format."""
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

FIGS = "/home/dheirav/pokerbot-scratch/ieee"
OUT = "/home/dheirav/pokerbot-scratch/NashForge_IEEE.docx"
FONT = "Times New Roman"

doc = Document()

# IEEE conference margins
base = doc.sections[0]
base.top_margin = Inches(0.75)
base.bottom_margin = Inches(1.0)
base.left_margin = base.right_margin = Inches(0.625)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X", "XI"]
_sec = [0]
_sub = [0]
_fig = [0]
_tab = [0]


def set_columns(section, num=2, space=288):
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space))
    cols.set(qn("w:equalWidth"), "1")


def rich(paragraph, text, size=10, italic=False):
    for piece in re.split(r"(\*\*.+?\*\*)", text):
        if not piece:
            continue
        run = paragraph.add_run(piece[2:-2] if piece.startswith("**") else piece)
        run.bold = piece.startswith("**")
        run.italic = italic
        run.font.name = FONT
        run.font.size = Pt(size)


def body(text, first_indent=True, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_indent:
        p.paragraph_format.first_line_indent = Inches(0.2)
    p.paragraph_format.space_after = Pt(0)
    rich(p, text, size=size)
    return p


def section(title):
    _sec[0] += 1
    _sub[0] = 0
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{ROMAN[_sec[0] - 1]}.  {title.upper()}")
    run.font.name = FONT
    run.font.size = Pt(10)
    run.small_caps = True
    return p


def subsection(title):
    _sub[0] += 1
    letter = chr(ord("A") + _sub[0] - 1)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"{letter}.  {title}")
    run.italic = True
    run.font.name = FONT
    run.font.size = Pt(10)
    return p


def figure(path, caption, width=3.30):
    _fig[0] += 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(8)
    rich(cap, f"Fig. {_fig[0]}.  {caption}", size=8)
    return cap


def table(caption, header, rows, size=7.5, full_width=False):
    _tab[0] += 1
    numeral = ROMAN[_tab[0] - 1]

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(f"TABLE {numeral}\n{caption.upper()}")
    run.font.name = FONT
    run.font.size = Pt(8)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(t.rows[0].cells, header):
        cell.text = ""
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.name = FONT
        r.font.size = Pt(size)
    for row in rows:
        for cell, text in zip(t.add_row().cells, row):
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(text))
            r.font.name = FONT
            r.font.size = Pt(size)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)
    return t


# ======================================================== title (1 column) ===
set_columns(base, 1)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run("NashForge: A Multi-Agent Framework for Comparative Evaluation of "
              "Game-Theoretic, Evolutionary, and Reinforcement Learning Agents in "
              "Imperfect-Information Games")
r.font.name = FONT; r.font.size = Pt(20)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
r = p.add_run("Dheirav"); r.font.name = FONT; r.font.size = Pt(11)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Roll No. 2023103027\nDepartment of Computer Science and Engineering")
r.font.name = FONT; r.font.size = Pt(9); r.italic = True

# ======================================================== body (2 columns) ===
two = doc.add_section(WD_SECTION.CONTINUOUS)
set_columns(two, 2)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Abstract—"); r.bold = True; r.italic = True; r.font.size = Pt(9); r.font.name = FONT
r = p.add_run(
    "Imperfect-information games such as poker require an agent to act without knowing the "
    "opponent's private state, and conventional approaches evaluate such agents against "
    "opponents drawn from their own population, producing systems that appear to improve while "
    "bearing no measurable relation to strong play. This paper presents NashForge, a framework "
    "for training and comparatively evaluating three families of agent — counterfactual regret "
    "minimization, evolutionary search, and proximal policy optimization — on a single verified "
    "poker engine using identical instrumentation. The framework validates its own measurement "
    "apparatus before reporting any result: the solver is checked against Kuhn poker's analytic "
    "game value of −1/18 and against exact exploitability on Leduc Hold'em, and every figure "
    "carries a confidence interval. The principal experimental finding is a crossover in card "
    "abstraction quality: a cheap made-hand bucketing outperforms Monte Carlo equity bucketing "
    "below approximately 640 seconds of training, after which equity bucketing wins by "
    "0.916 ± 0.118 chips per hand while using one fifth of the opponent's iterations. A "
    "secondary and negative finding is that the feature representation used by the learning "
    "agents recovers almost none of a hand's true value once its single card feature is removed, "
    "placing a ceiling on those agents that additional training cannot raise.")
r.bold = True; r.italic = True; r.font.size = Pt(9); r.font.name = FONT

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Index Terms—"); r.bold = True; r.italic = True; r.font.size = Pt(9); r.font.name = FONT
r = p.add_run("counterfactual regret minimization, imperfect-information games, card "
              "abstraction, exploitability, reinforcement learning, computer poker")
r.italic = True; r.font.size = Pt(9); r.font.name = FONT

# ------------------------------------------------------------ I. Intro ------
section("Introduction")
body("Imperfect-information games present a fundamentally harder computational problem than "
     "perfect-information games such as chess. A player must act without observing the "
     "opponent's private cards, and must therefore reason about a probability distribution over "
     "hidden states rather than about a single known position. This distinction removes the "
     "guarantees that make minimax search effective and requires an entirely different family of "
     "algorithms.", first_indent=False)
body("The dominant approach is counterfactual regret minimization [2], which decomposes the "
     "regret of a strategy into independent terms at each information set and converges toward a "
     "Nash equilibrium in two-player zero-sum games. Its practical application to no-limit "
     "Texas Hold'em requires two forms of approximation — grouping similar hands into buckets, "
     "and restricting the continuous space of bet sizes to a small discrete set — each of which "
     "introduces error that is rarely measured.")
body("This paper argues that the measurement apparatus deserves at least as much scrutiny as the "
     "learning algorithm. The framework described here was developed after an internal audit "
     "established that every result previously produced by its evolutionary training pipeline "
     "had been generated by a fitness function that scored the wrong player, and by a card "
     "dealer that re-dealt identical hands. An untrained random network scored +451 BB/100 under "
     "that metric, where a score of approximately zero is required by the zero-sum property of "
     "the game. All such results were withdrawn.")
body("The contributions of this work are threefold. First, a verified poker engine and solver "
     "validated against independently known answers. Second, an experimental determination of "
     "when a finer card abstraction repays its computational cost, which is shown to depend on "
     "the training budget rather than being a fixed property of the abstraction. Third, an "
     "empirical audit of the observation supplied to learning agents, which quantifies an "
     "information limit that no amount of training can overcome.")

# ---------------------------------------------------- II. Related Work ------
section("Related Work")

subsection("Theoretical Foundations")
body("Kuhn [1] introduced a three-card poker variant small enough to solve analytically, deriving "
     "an exact game value of −1/18 to the first player. Although trivial as poker, its closed-form "
     "solution makes it the standard first correctness test for any solver: an implementation "
     "that fails to reproduce −1/18 is provably incorrect regardless of how plausible its play "
     "appears. Zinkevich et al. [2] subsequently introduced counterfactual regret minimization, "
     "showing that overall regret in an extensive-form game decomposes into independent regret "
     "terms at each information set, and that the average strategy across iterations — not the "
     "current strategy — converges to equilibrium. This decomposition makes an otherwise "
     "intractable optimisation feasible and underpins essentially all later progress in the "
     "field.", first_indent=False)

subsection("Scaling Regret Minimization")
body("Vanilla CFR traverses the entire game tree per iteration, which is impossible at the scale "
     "of no-limit Hold'em. Lanctot et al. [3] introduced Monte Carlo sampling variants, of which "
     "external sampling is most widely used: chance outcomes and opponent actions are sampled "
     "while the traverser's own actions remain fully enumerated, preserving convergence in "
     "expectation while reducing per-iteration cost by orders of magnitude. Tammelin [4] proposed "
     "CFR+, which floors accumulated regrets at zero and weights the average strategy linearly, "
     "reporting order-of-magnitude faster convergence; the method was later used to essentially "
     "solve heads-up limit Hold'em. Brown and Sandholm [5] generalised this weighting through "
     "discounted and linear schedules, on the reasoning that iterations performed before the "
     "strategy stabilises carry misleading information. All four rules are implemented and "
     "compared here, and notably the published ordering did not reproduce under external "
     "sampling.", first_indent=False)

subsection("Abstraction")
body("Because no-limit Hold'em contains on the order of 10^160 information sets, hands must be "
     "grouped before any solver can run. Johanson et al. [6] systematically evaluated competing "
     "bucketing schemes and established clustering on hand-strength distributions as the "
     "strongest general approach, a method this work follows directly. Waugh et al. [7] reported "
     "the counter-intuitive result that refining an abstraction can produce a strategy that is "
     "more exploitable rather than less, demonstrating that the relationship between granularity "
     "and quality is not monotonic. That finding is the direct justification for the central "
     "experiment presented in Section VI. Ganzfried and Sandholm [8] addressed the complementary "
     "problem of responding to bet sizes outside the abstraction, showing that rounding to the "
     "nearest size is exploitable because it creates a deterministic boundary an opponent can sit "
     "just inside, and deriving the pseudo-harmonic mapping that this framework adopts.",
     first_indent=False)

subsection("Measuring Strategy Quality")
body("Southey et al. [9] introduced Leduc Hold'em, a six-card two-round variant sized so that "
     "exact best responses can be computed by exhaustive traversal while genuine bluffing "
     "structure remains. It has since become the standard intermediate benchmark and is used here "
     "as the largest game in which convergence can be verified against ground truth. For no-limit, "
     "where exact computation is impossible, Lisý and Bowling [10] proposed Local Best Response, "
     "a cheap adversary that maintains a belief over opponent holdings and greedily maximises "
     "immediate expected value. Because it explores only a subset of the strategies available to "
     "a true best response, its winnings constitute a rigorous lower bound: a large value proves "
     "weakness, while a small value proves nothing. Their striking result was that several "
     "published bots proved far more exploitable than head-to-head records suggested.",
     first_indent=False)

subsection("Landmark Systems")
body("DeepStack [11] was the first program to defeat professionals at heads-up no-limit, "
     "replacing precomputed abstraction with continual re-solving of local subgames and using "
     "neural networks to estimate values beyond its lookahead horizon. Libratus [12] defeated four "
     "top professionals over 120,000 hands using a coarse blueprint strategy, nested subgame "
     "solving, and a self-improvement module that repaired the abstraction regions opponents were "
     "exploiting — illustrating that a fixed abstraction is a persistent liability rather than a "
     "one-time simplification. Pluribus [13] extended these results to six-player poker, where "
     "equilibrium guarantees no longer apply, which motivates the deliberate ordering of heads-up "
     "before multi-player evaluation adopted in this work.", first_indent=False)

subsection("Learning-Based Approaches")
body("Proximal policy optimization [14] constrains each policy update through a clipped surrogate "
     "objective, preventing the destructive updates that destabilise simpler policy-gradient "
     "methods, and is adopted here as the reinforcement learning agent. Its known limitation is "
     "that self-play policy gradient does not converge to equilibrium in imperfect-information "
     "games. Neural Fictitious Self-Play [15] addresses this by combining reinforcement learning "
     "against the opponent's average strategy with supervised learning of the agent's own average "
     "behaviour. Deep CFR [16] replaces CFR's regret tables with neural networks, removing the "
     "need for hand-designed card abstraction entirely. Salimans et al. [17] demonstrated that "
     "population-based black-box optimisation can rival gradient methods while parallelising more "
     "cleanly, justifying the inclusion of an evolutionary agent — though it also implies the "
     "limitation expected here, that an optimiser receiving one scalar per episode performs no "
     "credit assignment across the decisions within it.", first_indent=False)

# ------------------------------------------------ III. System Architecture ---
section("System Architecture")
body("The system is organised as nine modules in a layered pipeline, each depending only on "
     "those beneath it. This separation is what permits three dissimilar learning algorithms to "
     "be compared fairly: they share the game, the abstraction and the evaluation harness, and "
     "differ only in the module that produces a strategy.", first_indent=False)

ARCH = [
    ("Poker Engine",
     "A complete Texas Hold'em implementation handling dealing, four betting rounds, "
     "minimum-raise enforcement, side pots and hand evaluation across eleven categories, built as "
     "an event-driven state machine in which each action transitions the game to a new legal "
     "state. Since every result is denominated in chips this module accounts for, it was verified "
     "by driving 2,400 randomised hands and confirming exact chip conservation, correct side-pot "
     "allocation, and agreement between two independent evaluators over 2,000 showdowns."),
    ("Traversable Game Interface",
     "Regret minimization traverses a game rather than playing it, recursing into all actions and "
     "backing up, which a conventional simulator cannot support because it mutates state in place. "
     "This module provides an immutable-state interface in which each state is a value, "
     "implemented for Kuhn poker, Leduc Hold'em and abstracted heads-up no-limit Hold'em."),
    ("Card Abstraction",
     "Two clustering strategies are used. Preflop, all 169 strategically distinct starting hands "
     "are enumerated exactly and grouped by the Chen formula. Postflop, where enumeration is "
     "impossible, situations are sampled and clustered by one-dimensional k-means. Two competing "
     "strength signals are implemented — Monte Carlo equity, which accounts for draws, and "
     "made-hand strength, which is far cheaper but blind to them."),
    ("Bet Abstraction",
     "A six-action space — fold, check/call, half-pot raise, pot raise, double-pot raise and "
     "all-in — bounds the continuous space of wagers. The resulting betting tree is enumerated "
     "combinatorially so that the size of the abstract game is computed exactly, and multiplying "
     "by the available buckets yields the memory requirement before training begins."),
    ("Regret Minimization Solvers",
     "Vanilla CFR is implemented for small games where full traversal is feasible, and "
     "external-sampling Monte Carlo CFR for no-limit. Four regret update rules are implemented "
     "and compared. The module accumulates counterfactual regret at each information set and "
     "derives the average strategy, which is the quantity that converges toward equilibrium."),
    ("Exploitability Measurement",
     "On small games exploitability is computed exactly by constructing a best response through "
     "tree traversal. On no-limit, Local Best Response provides a lower bound. The framework "
     "reports the one-sided nature of that guarantee explicitly rather than treating a small "
     "value as evidence of strength."),
    ("Action Translation",
     "The pseudo-harmonic mapping distributes an off-tree bet probabilistically across its two "
     "neighbouring abstract sizes. Nearest-neighbour rounding was rejected because its "
     "deterministic boundary permits an adversary to inflate measured exploitability with an "
     "artifact of the mapping rather than a property of the agent."),
    ("Evaluation Harness",
     "Agents are played against each other using duplicate play, a variance-reduction technique "
     "in which each hand is dealt twice with the seats swapped and identical cards replayed, so "
     "that position and card luck both cancel. Every result carries a standard error, and the "
     "harness declines to declare a winner when a difference lies inside its own noise."),
    ("Observation and Learning Layer",
     "The evolutionary and reinforcement learning agents consume a numerical feature vector — "
     "position, stack depth, pot odds, betting round and hand strength — rather than an abstract "
     "bucket. This vector feeds a policy network trained either by evolutionary selection and "
     "mutation or by proximal policy optimization. Its quality was audited experimentally, since "
     "a learning algorithm cannot recover information its inputs never contained."),
]
for name, text in ARCH:
    subsection(name)
    body(text, first_indent=False)

# ------------------------------------------------------- IV. Agents ---------
section("Agents and Implementation")
body("Six agent types are implemented against a common interface, so that any may be substituted "
     "into any matchup without modifying the harness. The CFR agent plays a strategy produced by "
     "regret minimization and is the only agent with a theoretical guarantee, serving as the "
     "reference point. The evolutionary agent is a policy network optimised by selection and "
     "Gaussian mutation with elite preservation. The PPO agent is a policy network trained from "
     "self-play experience using clipped updates and generalised advantage estimation. The Local "
     "Best Response agent is an adversary rather than a competitor, existing to attack a trained "
     "strategy and measure its exploitability. A random agent and an always-call agent provide "
     "reproducible baselines: the former is the floor any trained agent must clear, and the "
     "latter can only be exploited by an agent that bets for value.", first_indent=False)
body("The implementation is deliberately dependency-light, since a project whose central claim "
     "concerns measurement should not rest on components it cannot inspect. The engine, solvers, "
     "abstraction and harness are written from scratch in Python 3.12 using NumPy for vectorised "
     "arithmetic and Numba for just-in-time compilation of the hand evaluator, which lies on the "
     "hot path of every traversal. PyTorch provides the actor-critic network. The system carries "
     "197 regression tests, and every experimental result is stored as JSON alongside the script "
     "that produced it.")

# ------------------------------------------------ V. Experimental Setup -----
section("Experimental Methodology")
body("All experiments are conducted in heads-up no-limit Hold'em with 200-chip stacks and blinds "
     "of one and two chips. Card abstractions use six buckets per street and the betting "
     "abstraction permits one raise per street, which is the parameter that determines whether "
     "the abstract game fits in memory. Every reported figure is averaged over three random seeds "
     "and carries a standard error computed as the larger of the within-seed and between-seed "
     "estimates, so that a configuration which is merely fortunate in one seed cannot appear "
     "precise.", first_indent=False)
body("Matches use duplicate play with 20,000 hands per matchup and alternating seats. Training "
     "budgets are expressed in wall-clock seconds, because the practical question a practitioner "
     "faces is which abstraction to use given available time rather than given a fixed iteration "
     "count. Machine load is sampled during each training window and stored with every "
     "measurement, since a wall-clock budget measures the machine as much as the algorithm.")

# ---------------------------------------------------------- VI. Results -----
section("Results and Analysis")

subsection("Solver Validation")
body("Before any comparison is meaningful the solver must be shown to converge. Leduc Hold'em is "
     "small enough for exploitability to be computed exactly by traversal, so convergence is "
     "observed directly rather than inferred. Four update rules were run for 64,000 iterations "
     "across three seeds each; results appear in Fig. 1.", first_indent=False)
figure(f"{FIGS}/fig1_leduc.png",
       "Exact exploitability on Leduc Hold'em, logarithmic on both axes. All four update rules "
       "converge. CFR+ and DCFR finish worse than plain vanilla (0.0819 and 0.0930 against "
       "0.0615), contrary to published results [4], [5]; the likely cause is that those rules "
       "were designed for full traversal whereas this experiment uses external sampling.")

subsection("The Cost and Benefit of Card Abstraction")
body("Monte Carlo equity costs 44.05 ms per solver iteration against 8.11 ms for made-hand "
     "strength, a factor of 5.4, so the cheaper signal purchases roughly five times more "
     "training. Whether the additional iterations outweigh the lost information is not obvious a "
     "priori, and following the pathology reported by Waugh et al. [7] it cannot be assumed. It "
     "was resolved by playing the two resulting agents against each other across a ladder of "
     "training budgets, shown in Fig. 2 and Table I.", first_indent=False)
figure(f"{FIGS}/fig2_crossover.png",
       "Head-to-head result by training budget, in chips per hand to the equity agent. Bars right "
       "of zero indicate equity bucketing winning. Three seeds, 20,000 hands per matchup, with "
       "95% confidence intervals. A crossover exists between 640 and 2560 seconds.")
table("Head-to-head result with underlying iteration counts",
      ["Budget", "Equity it.", "Made-hand it.", "Chips/hand", "Verdict"],
      [["40 s", "1,075", "5,600", "−3.136 ± 0.729", "made-hand"],
       ["160 s", "3,808", "19,758", "−1.531 ± 0.164", "made-hand"],
       ["640 s", "14,850", "73,158", "+0.273 ± 0.203", "not sep."],
       ["2560 s", "60,117", "288,050", "+0.916 ± 0.118", "equity"]])
body("The cheap abstraction wins decisively at short budgets, where its fivefold iteration "
     "advantage dominates. The advantage vanishes near 640 seconds, and by 2560 seconds equity "
     "bucketing wins by 0.916 ± 0.118 chips per hand, approximately 7.8 standard errors from "
     "zero, while training on 60,117 iterations against its opponent's 288,050. All three seeds "
     "agree in direction at every budget.", first_indent=False)

subsection("Audit of the Learning Agents' Observation")
body("Because the evolutionary and reinforcement learning agents consume a feature vector rather "
     "than an abstract bucket, the quality of that vector bounds what either can learn. It was "
     "audited by sampling 1,500 decision states, computing each state's true equity by Monte "
     "Carlo simulation, and measuring by regression how much of that equity the vector recovers. "
     "Results appear in Fig. 3.", first_indent=False)
figure(f"{FIGS}/fig3_features.png",
       "Coefficient of determination against true equity. Removing a single feature — hand "
       "strength — collapses explanatory power to 0.013, showing that the remaining sixteen "
       "features carry almost no information about the player's own cards. Four inexpensive "
       "binary draw and texture features raise the postflop figure from 0.475 to 0.637.")
body("Principal component analysis established that the seventeen-slot vector carries only six "
     "independent dimensions. More consequentially, postflop states sharing a near-identical "
     "hand-strength value — states the agent receives as indistinguishable inputs — were measured "
     "to differ in true equity by as much as 0.907, meaning one hand may be a 5% underdog and "
     "another a 95% favourite while appearing identical to the network. This is an information "
     "limit rather than a modelling one, and no amount of additional training can overcome it.",
     first_indent=False)
table("Measured properties by module",
      ["Module", "Property measured", "Result"],
      [["Engine", "Chip conservation, 2,400 hands", "exact"],
       ["Engine", "Evaluator agreement, 2,000 showdowns", "0 disagreements"],
       ["Interface", "Kuhn value vs analytic −1/18", "reproduced"],
       ["Solvers", "Leduc exploitability at 64k it.", "0.0559–0.0930"],
       ["Card abstr.", "Cost per iteration", "44.05 / 8.11 ms"],
       ["Bet abstr.", "Information sets, 6 buckets", "32,800–410,000"],
       ["Harness", "Noise cut by duplicate play", "≈ 14%"],
       ["Observation", "Effective dimensionality of 17", "6"],
       ["Exploitability", "Bound on converged strategy", "not established"]])

subsection("Evolutionary Search, Measured")
body("""With the measurement apparatus validated, the second family of agent was trained and evaluated: a population of neural policies optimised by selection and Gaussian mutation over fifty generations, heads-up, on the rebuilt nineteen-feature observation. The training budget was set from measurement rather than convention — re-scoring the same population under different hands moves a genome by 136 BB/100 at 2,000 hands and 31 at 24,000, so 24,000 hands per genome was used.""", first_indent=False)
body("""The result is reported as an endpoint comparison rather than a learning curve, for reasons given below. The final genome and an untrained genome drawn from the same initial distribution were played against the identical panel at 40,000 hands, giving an uncertainty of about 14 BB/100.""")
figure(f"{FIGS}/fig4_evolution.png",
       "Evolutionary search before and after fifty generations. Endpoint comparison at 40,000 hands per matchup. The evolved agent improves decisively against a random opponent and not at all against either the calling station or the solver.")
table("Endpoint comparison, 40,000 hands per matchup",
      ["Opponent", "Untrained", "After 50 gen.", "Difference", "Verdict"],
      [['random', '−13.4 ± 14', '+192.7 ± 15', '+206.1 ± 39', 'improved'], ['always-call', '−8.1 ± 14', '+0.5 ± 4', '+8.6 ± 28', 'no change'], ['CFR agent', '−403.9 ± 13', '−370.1 ± 14', '+33.8 ± 37', 'no change']])
body("""Evolutionary search learned to exploit randomness rather than to play poker. It punishes an opponent that folds and raises at random by 206 BB/100, a difference of more than five standard errors. It gains nothing against a station that never folds, which can only be beaten by betting for value. And against the solver-derived agent it remains beaten by 370 BB/100 after fifty generations, having improved by nothing measurable. What it learned did not transfer, which is a more specific finding than the prediction that the method would simply plateau.""", first_indent=False)
body("""Only the independent panel makes this visible. Self-play fitness sat near zero throughout, as the zero-sum property of poker requires, and reveals nothing about whether anything was learned — which is precisely why the earlier pipeline could improve against itself for eighty-nine runs while getting worse at the game.""")

# --------------------------------------------------- VII. Observations ------
section("Discussion")
body("Several results warrant comment beyond their numerical value. The ordering of update rules "
     "in Fig. 1 contradicts published comparisons, and is reported rather than suppressed; the "
     "most plausible explanation is the interaction between those rules and external sampling "
     "rather than an implementation defect, since all four converge correctly.", first_indent=False)
body("Local Best Response failed to establish a bound on a well-trained strategy despite four "
     "defects being corrected and three successive valuation models being implemented. The honest "
     "interpretation is that a greedy one-step adversary genuinely cannot beat a converged "
     "strategy within this abstraction — a legitimate if weak result, given that the guarantee is "
     "one-sided, rather than an outstanding defect. It is reported as such because a measurement "
     "that quietly fails is more damaging than one that fails visibly.")
body("""Three results during this phase looked like findings and were not, and each was caught by a check already built into the framework. A per-generation panel curve measured over 3,000 hands appeared to rise from +111 to +214 BB/100; across eleven readings it scattered with a standard deviation of 56 against a measurement error of 57, and was indistinguishable from a constant. A benchmark row reported the evolved agent beating the solver by 60.8 BB/100, alongside a 74.3% lookup miss rate — the counter that exists so a benchmark which has quietly become a second random opponent shows up as a number. And the explanation first offered for that miss rate, that the solver covers only half its abstraction, was itself wrong: every lookup had found its key, and the harness was discarding valid entries because the solver stores one probability per legal action at a node rather than one per abstract action. A measurement that is too coarse or subtly wrong does not return no result; it returns a plausible one.""")
body("Duplicate play reduced the standard error by approximately 14%, considerably less than "
     "anticipated. Investigation showed that per-hand outcomes swing by an entire stack depending "
     "on whether an all-in occurred, and that this variation is driven by the actions taken "
     "rather than the cards dealt, so sharing cards cannot cancel it. The technique remains in "
     "use as it is free and strictly beneficial, but hand count remains the dominant lever.")

# ------------------------------------------ VIII. Comparative Analysis ------
section("Comparative Analysis")
table("NashForge compared with conventional approaches",
      ["Capability", "Heuristic", "Standard RL", "NashForge"],
      [["Strategy source", "Rules", "Learned", "Three compared"],
       ["Theoretical guarantee", "None", "None", "Equilibrium"],
       ["Engine correctness", "Assumed", "Assumed", "Verified"],
       ["Validated vs known answers", "None", "Rare", "Kuhn, Leduc"],
       ["Exploitability measured", "No", "No", "Exact + bound"],
       ["Card abstraction", "Fixed", "Implicit", "Two, compared"],
       ["Benchmark independence", "Self-play", "Self-play", "Unrelated panel"],
       ["Variance reduction", "None", "More hands", "Duplicate play"],
       ["Confidence intervals", "No", "Rare", "Every figure"],
       ["Observation audited", "No", "No", "Yes"]])

# --------------------------------------------------------- IX. Conclusion ---
section("Conclusion and Future Work")
body("This paper presented NashForge, a framework for comparatively evaluating game-theoretic, "
     "evolutionary and reinforcement learning agents in imperfect-information games under "
     "identical instrumentation. The framework's engine and solver are validated against "
     "independently known answers, its abstractions are measured rather than assumed, and every "
     "reported figure carries an explicit confidence interval.", first_indent=False)
body("The principal experimental result is that the value of a finer card abstraction depends on "
     "the training budget: made-hand bucketing wins below approximately 640 seconds and equity "
     "bucketing wins above it, at 7.8 standard errors and on one fifth of the opponent's "
     "iterations. The most instructive result, however, is negative — the observation audit "
     "demonstrated that the feature representation used by the learning agents discards most of "
     "the information distinguishing a strong hand from a weak one, imposing a ceiling that "
     "further training cannot raise. Identifying such a limit before investing in training is "
     "precisely what a measurement-first framework is intended to deliver.")
body("Future work follows three directions. The observation layer will be rebuilt to include "
     "draw and board-texture features, which the audit shows raises postflop explanatory power by "
     "34%. The evolutionary and PPO agents will then be trained and measured against the "
     "solver-derived reference. Finally, extending Local Best Response with multi-street "
     "lookahead may yield a bound that the present one-step formulation cannot.")

# --------------------------------------------------------- References -------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(11)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("REFERENCES"); r.font.name = FONT; r.font.size = Pt(10); r.small_caps = True

REFS = [
 'H. W. Kuhn, "Simplified two-person poker," in Contributions to the Theory of Games, vol. 1, Princeton, NJ, USA: Princeton Univ. Press, 1950, pp. 97–103.',
 'M. Zinkevich, M. Johanson, M. Bowling, and C. Piccione, "Regret minimization in games with incomplete information," in Proc. Adv. Neural Inf. Process. Syst. (NIPS), 2007.',
 'M. Lanctot, K. Waugh, M. Zinkevich, and M. Bowling, "Monte Carlo sampling for regret minimization in extensive games," in Proc. Adv. Neural Inf. Process. Syst. (NIPS), 2009.',
 'O. Tammelin, "Solving large imperfect information games using CFR+," arXiv:1407.5042, 2014.',
 'N. Brown and T. Sandholm, "Solving imperfect-information games via discounted regret minimization," in Proc. AAAI Conf. Artif. Intell., 2019.',
 'M. Johanson, N. Burch, R. Valenzano, and M. Bowling, "Evaluating state-space abstractions in extensive-form games," in Proc. Int. Conf. Auton. Agents Multi-Agent Syst. (AAMAS), 2013.',
 'K. Waugh, D. Schnizlein, M. Bowling, and D. Szafron, "Abstraction pathologies in extensive games," in Proc. Int. Conf. Auton. Agents Multi-Agent Syst. (AAMAS), 2009.',
 'S. Ganzfried and T. Sandholm, "Action translation in extensive-form games with large action spaces," in Proc. Int. Joint Conf. Artif. Intell. (IJCAI), 2013.',
 'F. Southey, M. Bowling, B. Larson, C. Piccione, N. Burch, D. Billings, and C. Rayner, "Bayes\' bluff: Opponent modelling in poker," in Proc. Conf. Uncertainty Artif. Intell. (UAI), 2005.',
 'V. Lisý and M. Bowling, "Equilibrium approximation quality of current no-limit poker bots," arXiv:1612.07547, 2016.',
 'M. Moravčík, M. Schmid, N. Burch, V. Lisý, D. Morrill, N. Bard, T. Davis, K. Waugh, M. Johanson, and M. Bowling, "DeepStack: Expert-level artificial intelligence in heads-up no-limit poker," Science, vol. 356, no. 6337, pp. 508–513, 2017.',
 'N. Brown and T. Sandholm, "Superhuman AI for heads-up no-limit poker: Libratus beats top professionals," Science, vol. 359, no. 6374, pp. 418–424, 2018.',
 'N. Brown and T. Sandholm, "Superhuman AI for multiplayer poker," Science, vol. 365, no. 6456, pp. 885–890, 2019.',
 'J. Schulman, F. Wolski, P. Dhariwal, A. Radford, and O. Klimov, "Proximal policy optimization algorithms," arXiv:1707.06347, 2017.',
 'J. Heinrich and D. Silver, "Deep reinforcement learning from self-play in imperfect-information games," arXiv:1603.01121, 2016.',
 'N. Brown, A. Lerer, S. Gross, and T. Sandholm, "Deep counterfactual regret minimization," in Proc. Int. Conf. Mach. Learn. (ICML), 2019.',
 'T. Salimans, J. Ho, X. Chen, S. Sidor, and I. Sutskever, "Evolution strategies as a scalable alternative to reinforcement learning," arXiv:1703.03864, 2017.',
]
for index, ref in enumerate(REFS, 1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.22)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f"[{index}]  {ref}")
    r.font.name = FONT
    r.font.size = Pt(8)

doc.save(OUT)
print(f"wrote {OUT}")
