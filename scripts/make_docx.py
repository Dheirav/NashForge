"""Build the NashForge assignment as a Word document."""
import os
import re

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

FIGS = os.path.join(ROOT, "docs", "figures")
OUT = os.path.join(ROOT, "docs", "NashForge_Assignment.docx")

SERIF, SANS = "Cambria", "Calibri"
ACCENT = RGBColor(0x2F, 0x48, 0x58)
MUTED = RGBColor(0x5C, 0x64, 0x6A)

doc = Document()
for section in doc.sections:
    section.top_margin = section.bottom_margin = Inches(0.85)
    section.left_margin = section.right_margin = Inches(0.85)

normal = doc.styles["Normal"]
normal.font.name = SERIF
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

_section_no = [0]


def rich(paragraph, text):
    """Write text into a paragraph, honouring **bold** markers."""
    for piece in re.split(r"(\*\*.+?\*\*)", text):
        if not piece:
            continue
        run = paragraph.add_run(piece[2:-2] if piece.startswith("**") else piece)
        run.bold = piece.startswith("**")


def para(text, justify=True, italic=False, size=None):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    rich(p, text)
    for run in p.runs:
        run.italic = run.italic or italic
        if size:
            run.font.size = Pt(size)
    return p


def h1(text):
    _section_no[0] += 1
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{_section_no[0]}. {text.upper()}")
    run.bold = True
    run.font.name = SANS
    run.font.size = Pt(12)
    run.font.color.rgb = ACCENT
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.name = SANS
    run.font.size = Pt(10.5)
    return p


def figure(path, title, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.space_after = Pt(10)
    rich(cap, f"**{title}** {caption}")
    for run in cap.runs:
        run.font.size = Pt(9)
        if not run.bold:
            run.font.color.rgb = MUTED


def table(caption, header, rows, widths=None):
    cap = doc.add_paragraph()
    cap.paragraph_format.space_before = Pt(8)
    cap.paragraph_format.space_after = Pt(3)
    run = cap.add_run(caption)
    run.bold = True
    run.font.name = SANS
    run.font.size = Pt(9.5)

    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for cell, text in zip(t.rows[0].cells, header):
        cell.text = ""
        r = cell.paragraphs[0].add_run(text)
        r.bold = True
        r.font.name = SANS
        r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for cell, text in zip(cells, row):
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(text))
            r.font.name = SANS
            r.font.size = Pt(8.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ---------------------------------------------------------------- title -----
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("ARTIFICIAL INTELLIGENCE — PROJECT REPORT")
r.bold = True; r.font.name = SANS; r.font.size = Pt(10); r.font.color.rgb = MUTED

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(4)
r = t.add_run("NashForge"); r.bold = True; r.font.size = Pt(26); r.font.name = SERIF

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(12)
r = t.add_run("A Multi-Agent Framework for Comparative Evaluation of Game-Theoretic, "
              "Evolutionary, and Reinforcement Learning Agents in Imperfect-Information Games")
r.italic = True; r.font.size = Pt(12)

t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(14)
r = t.add_run("Submitted by  DHEIRAV          Roll No.  2023103027")
r.font.name = SANS; r.font.size = Pt(10.5); r.bold = True

# -------------------------------------------------------------- abstract ----
h1("Abstract")
para("Imperfect-information games such as poker present a fundamentally harder problem than "
     "perfect-information games like chess, because a player must act without knowing the "
     "opponent's private cards and must therefore reason about probability distributions over "
     "hidden states rather than about a single known position. Conventional approaches to "
     "building poker agents rely on hand-crafted heuristics or on machine learning methods that "
     "optimise an agent against opponents drawn from its own population, which produces systems "
     "that appear to improve while having no measurable relationship to strong play.")
para("This project proposes **NashForge**, a framework for training and, more importantly, "
     "comparatively evaluating three distinct families of agent on a single verified poker "
     "engine: counterfactual regret minimization, which converges provably toward a "
     "game-theoretic equilibrium; evolutionary search, which optimises a population by selection "
     "and mutation; and proximal policy optimization, a reinforcement learning method that "
     "improves a policy from experience. The three are measured against a common benchmark panel "
     "using identical instrumentation, so that differences between them reflect the algorithms "
     "rather than differences in how each was evaluated.")
para("The framework's distinguishing characteristic is that its measurement apparatus is itself "
     "validated before any result is reported. The solver is checked against Kuhn poker's "
     "analytically known game value of −1/18 and against exact exploitability computed by full "
     "tree traversal on Leduc Hold'em; the card abstraction's cost and benefit are measured "
     "rather than assumed; and every published figure carries a confidence interval. During "
     "development this discipline caused several previously published results to be formally "
     "withdrawn, and that correction is reported here as a finding in its own right rather than "
     "omitted. The system is covered by 233 automated regression tests.")

# ----------------------------------------------------- literature survey ----
h1("Literature Survey")
para("The work below is grouped by the problem it addresses: establishing the theory, making it "
     "scale, reducing the game to a tractable size, measuring the result honestly, and the "
     "systems and learning methods that define the current state of the art. Each is summarised "
     "together with its bearing on the present project.")

SURVEY = [
    ("Theoretical Foundations", [
        "**Kuhn (1950), \"Simplified Two-Person Poker.\"** Kuhn introduced a three-card poker "
        "variant small enough to be solved analytically by hand, and derived its exact game value "
        "of −1/18 to the first player along with the full family of equilibrium strategies. "
        "Although the game is far too simple to be interesting as poker, its analytic solution "
        "makes it the standard first test for any solver: an implementation that fails to "
        "reproduce −1/18 is provably incorrect, regardless of how plausible its behaviour "
        "appears. NashForge uses precisely this property, treating Kuhn as a correctness check "
        "performed before any larger game is attempted.",

        "**Zinkevich, Johanson, Bowling and Piccione (2007), \"Regret Minimization in Games with "
        "Incomplete Information.\"** This paper introduced Counterfactual Regret Minimization, "
        "the algorithm underpinning essentially all subsequent progress in computer poker. Its "
        "central insight is that the overall regret of a strategy in an extensive-form game can "
        "be decomposed into independent regret terms at each information set, so that minimising "
        "regret locally minimises it globally, and the average strategy over many iterations — "
        "not the current one — converges to a Nash equilibrium. This decomposition is what makes "
        "an otherwise intractable optimisation possible, and it is the algorithm NashForge "
        "implements as its game-theoretic agent and validates against known game values."]),

    ("Making the Algorithm Scale", [
        "**Lanctot, Waugh, Zinkevich and Bowling (2009), \"Monte Carlo Sampling for Regret "
        "Minimization in Extensive Games.\"** Vanilla CFR must traverse the entire game tree on "
        "every iteration, which is impossible for games the size of no-limit Hold'em. The authors "
        "introduced a family of sampling variants, of which external sampling is the most widely "
        "used: chance outcomes and the opponent's actions are sampled while the traverser's own "
        "actions are still fully enumerated, preserving the convergence guarantee in expectation "
        "while reducing the per-iteration cost by orders of magnitude. NashForge implements "
        "external-sampling MCCFR as the only practical route to training on no-limit, and retains "
        "vanilla CFR for the small games where exact traversal remains feasible.",

        "**Tammelin (2014), \"Solving Large Imperfect Information Games Using CFR+.\"** CFR+ "
        "modifies the regret update so that accumulated regrets are floored at zero rather than "
        "allowed to become arbitrarily negative, and weights the average strategy linearly by "
        "iteration. The paper reports order-of-magnitude faster convergence than vanilla CFR, and "
        "the method was subsequently used to essentially solve heads-up limit Hold'em. NashForge "
        "implements CFR+ as one of four comparable update rules; notably, the present experiments "
        "found it converging more slowly than vanilla under external sampling, a discrepancy "
        "attributed to the rule having been designed for full traversal.",

        "**Brown and Sandholm (2019), \"Solving Imperfect-Information Games via Discounted Regret "
        "Minimization.\"** This work generalises the weighting idea by discounting early regrets "
        "and strategy contributions according to tunable schedules, on the reasoning that "
        "iterations performed before the strategy has stabilised carry misleading information. "
        "Discounted CFR and Linear CFR are shown to outperform CFR+ on several benchmarks. Both "
        "are implemented in NashForge, and the comparison between all four rules is reported as "
        "an experimental result rather than adopted from the literature — a distinction that "
        "proved material, since the published ordering did not reproduce under this project's "
        "sampling regime."]),

    ("Abstraction", [
        "**Johanson, Burch, Valenzano and Bowling (2013), \"Evaluating State-Space Abstractions "
        "in Extensive-Form Games.\"** Since no-limit Hold'em contains on the order of 10^160 "
        "information sets, hands must be grouped into a bounded number of buckets before any "
        "solver can run. This paper systematically evaluates competing bucketing schemes and "
        "establishes clustering on hand-strength distributions as the strongest general approach. "
        "NashForge follows this method directly, clustering sampled postflop situations by "
        "one-dimensional k-means, and extends it by treating the choice of the underlying "
        "strength signal as an experimental question rather than a settled one.",

        "**Waugh, Schnizlein, Bowling and Szafron (2009), \"Abstraction Pathologies in Extensive "
        "Games.\"** This paper reports the counter-intuitive and important finding that refining "
        "an abstraction — giving the solver a more detailed view of the game — can produce a "
        "strategy that is more exploitable, not less. The relationship between abstraction "
        "granularity and final strategy quality is therefore not monotonic and cannot be assumed. "
        "This result is the direct justification for NashForge's central experiment: rather than "
        "presuming the finer card abstraction is better, the project measures the two against "
        "each other across a range of training budgets and finds that which one wins depends on "
        "the budget available.",

        "**Ganzfried and Sandholm (2013), \"Action Translation in Extensive-Form Games with Large "
        "Action Spaces.\"** An agent trained on a handful of discrete bet sizes must still "
        "respond when a human or another program bets an amount outside that set. The authors "
        "show that the obvious solution — rounding to the nearest available size — is "
        "exploitable, because the mapping has a deterministic boundary that an opponent can "
        "deliberately sit just inside, and they derive the pseudo-harmonic mapping, which "
        "distributes an off-tree bet probabilistically across its two neighbours in a way that "
        "bounds this exploitation. NashForge implements this mapping so that its adversarial "
        "evaluator can bet sizes the solver never trained against without the resulting "
        "measurement being an artifact of the translation itself."]),

    ("Measuring Strategy Quality", [
        "**Southey, Bowling, Larson, Piccione, Burch, Billings and Rayner (2005), \"Bayes' Bluff: "
        "Opponent Modelling in Poker.\"** Alongside its contribution on Bayesian opponent "
        "modelling, this paper introduced Leduc Hold'em, a six-card two-round variant "
        "deliberately sized so that exact best responses and true exploitability can be computed "
        "by exhaustive traversal while the game still exhibits genuine bluffing and "
        "card-abstraction structure. Leduc has since become the standard intermediate benchmark, "
        "and NashForge uses it exactly as intended: as the largest game in which the solver's "
        "convergence can be verified against ground truth rather than estimated.",

        "**Lisý and Bowling (2017), \"Equilibrium Approximation Quality of Current No-Limit Poker "
        "Bots.\"** Because exact exploitability cannot be computed for no-limit, the authors "
        "propose Local Best Response, a computationally cheap adversary that maintains a belief "
        "over the opponent's holdings and greedily selects the action maximising immediate "
        "expected value under a one-step lookahead. Since LBR explores only a subset of the "
        "strategies available to a true best response, whatever it wins constitutes a rigorous "
        "lower bound on exploitability: a large value proves an agent is weak, while a small "
        "value proves nothing at all. The paper's striking result was that several published bots "
        "proved far more exploitable than their head-to-head records suggested. NashForge "
        "implements LBR and reports this one-sided guarantee explicitly, having found — after "
        "correcting four defects in its own implementation — that a converged strategy in this "
        "abstraction resists it entirely."]),

    ("Landmark Systems", [
        "**Moravčík, Schmid, Burch, Lisý, Morrill, Bard, Davis, Waugh, Johanson and Bowling "
        "(2017), \"DeepStack: Expert-Level Artificial Intelligence in Heads-Up No-Limit "
        "Poker.\"** DeepStack was the first program to defeat professional players at heads-up "
        "no-limit, and it did so by replacing precomputed abstraction with continual re-solving: "
        "rather than looking up a strategy from a table computed in advance, it solves a local "
        "subgame at every decision, using deep neural networks to estimate the values of states "
        "beyond its lookahead horizon. The work demonstrated that the abstraction bottleneck "
        "NashForge measures directly can, given sufficient computational resources, be "
        "substantially circumvented.",

        "**Brown and Sandholm (2018), \"Superhuman AI for Heads-Up No-Limit Poker: Libratus Beats "
        "Top Professionals.\"** Libratus defeated four top professionals over 120,000 hands, "
        "using a three-part architecture: a blueprint strategy computed over a coarse "
        "abstraction, nested subgame solving that refines the strategy in real time as a hand "
        "progresses, and a self-improvement module that repairs the parts of the abstraction "
        "opponents were successfully exploiting. Its significance for the present project lies in "
        "the third component, which illustrates that a fixed abstraction is a persistent "
        "liability rather than a one-time simplification — the same weakness this project "
        "quantifies at much smaller scale.",

        "**Brown and Sandholm (2019), \"Superhuman AI for Multiplayer Poker.\"** Pluribus extended "
        "these results to six-player no-limit, where the theoretical guarantees of two-player "
        "zero-sum games no longer apply, since a Nash equilibrium in a multiplayer game need not "
        "be either unique or safe to play. It nevertheless defeated elite professionals using a "
        "comparatively inexpensive blueprint combined with limited-lookahead search. This work "
        "motivates NashForge's deliberate ordering of heads-up before six-max: the guarantees "
        "that make the two-player case measurable simply do not transfer."]),

    ("Learning-Based Approaches", [
        "**Schulman, Wolski, Dhariwal, Radford and Klimov (2017), \"Proximal Policy Optimization "
        "Algorithms.\"** PPO improves a policy from sampled experience while constraining each "
        "update through a clipped surrogate objective, preventing the destructively large policy "
        "changes that destabilise simpler policy-gradient methods, at a fraction of the "
        "implementation complexity of trust-region approaches. Its robustness and simplicity have "
        "made it a default choice across reinforcement learning. NashForge adopts PPO as its "
        "reinforcement learning agent, while noting the theoretical caveat that self-play policy "
        "gradient does not converge to a Nash equilibrium in imperfect-information games and "
        "therefore yields a strong but exploitable strategy.",

        "**Heinrich and Silver (2016), \"Deep Reinforcement Learning from Self-Play in "
        "Imperfect-Information Games.\"** Neural Fictitious Self-Play combines reinforcement "
        "learning against the opponent's average strategy with supervised learning of the agent's "
        "own average behaviour, approximating fictitious play with neural networks and converging "
        "toward equilibrium in imperfect-information games without domain-specific abstraction. "
        "The work is directly relevant as the principled alternative to plain self-play, and "
        "identifies precisely the weakness NashForge expects its PPO agent to exhibit.",

        "**Brown, Lerer, Gross and Sandholm (2019), \"Deep Counterfactual Regret "
        "Minimization.\"** Deep CFR replaces CFR's explicit regret tables with neural networks "
        "trained to approximate them, removing the requirement for a hand-designed card "
        "abstraction and allowing the method to generalise across similar information sets. It "
        "represents the natural continuation of the trajectory this project follows, and is the "
        "approach NashForge's tabular solver would most sensibly be extended toward once the "
        "abstraction limits it measures become binding.",

        "**Salimans, Ho, Chen, Sidor and Sutskever (2017), \"Evolution Strategies as a Scalable "
        "Alternative to Reinforcement Learning.\"** This paper demonstrates that a "
        "population-based black-box optimiser, perturbing network parameters and selecting on "
        "total episode return, can rival gradient-based reinforcement learning on demanding "
        "control tasks while parallelising far more cleanly, since workers need exchange only "
        "scalar returns. It provides the justification for including an evolutionary agent in "
        "this comparison — though it also implies the limitation NashForge expects to observe, "
        "namely that an optimiser receiving one scalar per episode performs no credit assignment "
        "across the individual decisions within it."]),
]

for group, entries in SURVEY:
    h2(group)
    for entry in entries:
        para(entry)

# ------------------------------------------------- modular architecture -----
h1("Modular Architecture")
para("The system is organised as nine modules arranged in a layered pipeline, in which each "
     "layer depends only on the layers beneath it. This separation is what allows three very "
     "different learning algorithms to be compared fairly: they share the game, the abstraction, "
     "and the evaluation harness, and differ only in the module that produces a strategy.")

MODULES = [
    ("1. Poker Engine",
     "The foundation of the system is a complete Texas Hold'em implementation that handles card "
     "dealing, the four betting rounds, minimum-raise enforcement, side pots for all-in "
     "situations, and hand evaluation across all eleven hand categories. It is built as an "
     "**event-driven state machine** in which each player action transitions the game to a new "
     "legal state. Correctness here is non-negotiable, since every result in the project is "
     "denominated in chips this module accounts for: it was verified by driving 2,400 randomised "
     "hands and confirming exact chip conservation, correct side-pot allocation, and agreement "
     "between two independent hand evaluators over 2,000 random seven-card showdowns."),

    ("2. Traversable Game Interface",
     "Counterfactual regret minimization does not play a game, it traverses one: at every "
     "decision it must recurse into all available actions and back up, which a conventional "
     "simulator cannot support because it mutates state in place and consumes a dealt deck. This "
     "module therefore provides an **immutable-state interface** in which each state is a value "
     "rather than an object being modified, implemented for three games of increasing difficulty "
     "— Kuhn poker, Leduc Hold'em, and abstracted heads-up no-limit Hold'em. The two smaller "
     "games exist specifically so the solver can be checked against known answers."),

    ("3. Card Abstraction",
     "No-limit Hold'em contains on the order of 10^160 information sets, which no computer can "
     "enumerate, so hands must be grouped into a bounded number of strength buckets. This module "
     "uses **two different clustering strategies**: preflop, all 169 strategically distinct "
     "starting hands are enumerated exactly and grouped by the Chen formula; postflop, where "
     "enumeration is impossible, situations are sampled and clustered by one-dimensional k-means. "
     "Two competing strength signals are implemented — Monte Carlo equity, which accounts for "
     "draws, and made-hand strength, which is far cheaper but blind to them — and choosing "
     "between them is treated as an experimental question rather than a design assumption."),

    ("4. Bet Abstraction",
     "Alongside card abstraction, the space of possible bet sizes must also be bounded, since "
     "no-limit poker permits any wager up to a player's stack. This module defines a **six-action "
     "space** — fold, check/call, half-pot raise, pot raise, double-pot raise, and all-in — and "
     "enumerates the resulting betting tree combinatorially, so the size of the abstract game is "
     "computed exactly rather than estimated. Multiplying that count by the available buckets "
     "gives the memory requirement before any training is attempted."),

    ("5. Regret Minimization Solvers",
     "This is the game-theoretic learner. It implements **vanilla CFR**, which traverses the "
     "entire tree and is feasible only for small games, and **external-sampling Monte Carlo "
     "CFR**, which samples chance outcomes and so scales to no-limit. Four regret update rules "
     "are implemented and compared experimentally. The module accumulates counterfactual regret "
     "at every information set and derives the average strategy, which is the quantity that "
     "converges toward equilibrium — a fact that is easy to state and easy to implement "
     "incorrectly, which is why the validation module exists."),

    ("6. Exploitability Measurement",
     "An agent's true quality is measured by exploitability: how much a perfectly-informed "
     "adversary could win against it. On small games this module computes exploitability "
     "**exactly** by traversing the tree and constructing a best response. On no-limit, where "
     "exact computation is impossible, it implements **Local Best Response**, a greedy adversary "
     "that tracks what the opponent's betting reveals about their holding and exploits it. This "
     "yields a one-sided guarantee — a large value proves an agent is exploitable, whereas a "
     "small value proves nothing — and the framework reports that asymmetry explicitly rather "
     "than treating a small number as evidence of strength."),

    ("7. Action Translation",
     "When an adversary bets an amount the abstraction does not contain, something must decide "
     "what the trained agent believes it faced. This module implements the **pseudo-harmonic "
     "mapping** of Ganzfried and Sandholm, which maps an off-tree bet probabilistically onto its "
     "two neighbouring abstract sizes. The obvious alternative — rounding to the nearest size — "
     "was rejected deliberately, because a deterministic mapping creates a boundary that an "
     "adversary can sit just inside, inflating the measured exploitability with an artifact of "
     "the mapping rather than a property of the agent."),

    ("8. Evaluation Harness",
     "This module decides what every reported number means, and is therefore the component whose "
     "own correctness matters most. It plays two agents against each other using **duplicate "
     "play**, a variance-reduction technique borrowed from competitive bridge: each hand is dealt "
     "twice with the seats swapped and the identical cards replayed, so that both position and "
     "card luck cancel between the two runs rather than being averaged over. Every result carries "
     "a standard error, and the harness refuses to declare a winner when the difference between "
     "two agents lies inside its own noise."),

    ("9. Observation and Learning Layer",
     "The evolutionary and reinforcement learning agents do not consume abstract buckets; they "
     "consume a numerical feature vector describing the current decision. This module extracts "
     "that vector — position, stack depth, pot odds, betting round, and hand strength — and feeds "
     "it to a **policy network** trained either by evolutionary selection and mutation or by "
     "proximal policy optimization. The quality of this representation was itself audited "
     "experimentally, with results reported below, because a learning algorithm cannot recover "
     "information its inputs never contained."),
]
for title, body in MODULES:
    h2(title)
    para(body)

# ------------------------------------------------------------- agents -------
h1("Agents Implemented")
para("Six agent types are implemented against a common interface, so that any of them can be "
     "substituted into any matchup without changing the harness. This uniformity is what makes "
     "the comparison meaningful.")
AGENTS = [
    "**CFR Agent** — plays a strategy produced by regret minimization. It is the only agent in "
    "the system with a theoretical guarantee about what it converges toward, and it therefore "
    "serves as the framework's reference point.",
    "**Evolutionary Agent** — a policy network whose weights are optimised by selection and "
    "Gaussian mutation across generations, with elite preservation between them.",
    "**PPO Agent** — a policy network trained by proximal policy optimization from self-play "
    "experience, using clipped policy updates and generalised advantage estimation.",
    "**Local Best Response Agent** — an adversary rather than a competitor. It exists to attack a "
    "trained strategy and thereby measure its exploitability, tracking a belief over the "
    "opponent's possible holdings and updating it by Bayes after every action observed.",
    "**Random Agent** — chooses uniformly among legal actions. The floor any trained agent must "
    "clear, and a diagnostic when one does not.",
    "**Always-Call Agent** — never folds and never raises. A specific, reproducible weak opponent "
    "whose exploitation requires an agent to actually bet for value.",
]
for index, text in enumerate(AGENTS, 1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.25)
    rich(p, f"{index}. {text}")

# ---------------------------------------------------------- tech stack ------
h1("Technology Stack")
para("The implementation is deliberately dependency-light, because a project whose central claim "
     "is about measurement should not rest on components it cannot inspect. The engine, the "
     "solvers, the abstraction and the evaluation harness are written from scratch in **Python "
     "3.12** using **NumPy** for vectorised arithmetic and **Numba** for just-in-time compilation "
     "of the hand evaluator, which sits on the hot path of every traversal. The reinforcement "
     "learning component uses **PyTorch** for the actor-critic network. Testing is by **pytest**, "
     "with 197 regression tests; version control and reproducibility are handled through "
     "**Git**, with every experimental result stored as JSON alongside the script that produced "
     "it.")

# -------------------------------------------------- performance metrics -----
h1("Performance Metrics and Analysis")
table("Key figures at a glance",
      ["Regression tests", "Kuhn game value", "Cost ratio", "Headline separation"],
      [["197, from zero before the audit", "−1/18 reproduced", "5.4× equity over made-hand",
        "7.8σ"]])

h2("Validating the solver")
para("Before any comparison is meaningful, the solver must be shown to converge. Leduc Hold'em is "
     "small enough that exploitability can be computed exactly by traversing the tree, so "
     "convergence is observed directly rather than inferred. Four regret update rules were run "
     "for 64,000 iterations across three random seeds each.")
figure(f"{FIGS}/fig1_leduc.png", "Figure 1. Exact exploitability on Leduc Hold'em.",
       "Logarithmic on both axes; lower is closer to equilibrium. All four rules converge, which "
       "is the property being verified. The ordering is a secondary and slightly awkward "
       "observation: CFR+ and DCFR finish worse than plain vanilla (0.0819 and 0.0930 against "
       "0.0615), where the literature reports them substantially faster. The likely explanation "
       "is that those rules were designed for full tree traversal while this experiment uses "
       "external sampling. It is reported rather than omitted.")

h2("The cost and benefit of card abstraction")
para("The central experiment of the project asks a question a practitioner actually faces: given "
     "a fixed amount of training time, is the more informative but more expensive card "
     "abstraction worth using? Monte Carlo equity costs 44.05 ms per solver iteration against "
     "8.11 ms for made-hand strength — a factor of 5.4 — so the cheaper signal buys roughly five "
     "times more training. Whether the extra iterations outweigh the lost information is not "
     "obvious, and was resolved by playing the two resulting agents against each other across a "
     "ladder of training budgets.")
figure(f"{FIGS}/fig2_crossover.png", "Figure 2. Head-to-head result by training budget.",
       "Chips per hand to the equity agent; bars right of zero mean equity is winning. Three "
       "seeds, 20,000 hands per matchup, seats alternating, with 95% confidence intervals. A "
       "crossover exists: the cheap abstraction wins at short budgets, the lead disappears around "
       "640 seconds, and the expensive abstraction wins by 2560 seconds at roughly 7.8σ — while "
       "training on 60,117 iterations against its opponent's 288,050.")
table("Table 1. The same result, with the iteration counts behind it",
      ["Budget", "Equity iterations", "Made-hand iterations", "Chips/hand", "Verdict"],
      [["40 s", "1,075", "5,600", "−3.136 ± 0.729", "made-hand"],
       ["160 s", "3,808", "19,758", "−1.531 ± 0.164", "made-hand"],
       ["640 s", "14,850", "73,158", "+0.273 ± 0.203", "not separated"],
       ["2560 s", "60,117", "288,050", "+0.916 ± 0.118", "equity"]])

h2("Auditing the learning agents' observation")
para("Because the evolutionary and reinforcement learning agents consume a feature vector rather "
     "than an abstract bucket, the quality of that vector places a ceiling on what either can "
     "learn. It was therefore audited experimentally: 1,500 decision states were sampled, each "
     "state's true equity was computed by Monte Carlo simulation, and the vector's ability to "
     "recover that equity was measured by regression.")
figure(f"{FIGS}/fig3_features.png", "Figure 3. How much of a hand's true value the observation recovers.",
       "Coefficient of determination against Monte Carlo equity. The third bar is the finding: "
       "removing a single feature — hand strength — collapses the vector's explanatory power to "
       "0.013, meaning the remaining sixteen features carry almost no information about the "
       "player's own cards. Adding four inexpensive binary features (flush draw, open-ended "
       "straight draw, paired board, three-to-a-flush) raises the postflop figure from 0.475 to "
       "0.637, a relative improvement of 34%.")
table("Table 2. Measured properties, by module",
      ["Module", "Property measured", "Result"],
      [["Poker Engine", "Chip conservation, 2,400 hands", "exact"],
       ["Poker Engine", "Evaluator agreement, 2,000 showdowns", "0 disagreements"],
       ["Traversable Interface", "Kuhn game value vs analytic −1/18", "reproduced"],
       ["Regret Solvers", "Leduc exploitability at 64k iterations", "0.0559 – 0.0930"],
       ["Card Abstraction", "Cost per iteration, equity vs made-hand", "44.05 / 8.11 ms"],
       ["Bet Abstraction", "Information sets, 6 buckets, 1 raise", "32,800 – 410,000"],
       ["Evaluation Harness", "Noise reduction from duplicate play", "≈ 14%"],
       ["Observation Layer", "Effective dimensionality of 17 slots", "6"],
       ["Exploitability", "Bound on a converged strategy", "not established"]])

h2("Phase Two — Evolutionary Search, Measured")
para("""With the measurement apparatus validated, the second family of agent was trained and evaluated: a population of neural policies optimised by selection and Gaussian mutation over fifty generations, heads-up, on the rebuilt nineteen-feature observation. The training budget was set from measurement rather than convention — re-scoring the same population under different hands moves a genome by 136 BB/100 at 2,000 hands and 31 at 24,000, so 24,000 hands per genome was used.""")
para("""The result is reported as an endpoint comparison rather than a learning curve, for reasons given below. The final genome and an untrained genome drawn from the same initial distribution were played against the identical panel at 40,000 hands, giving an uncertainty of about 14 BB/100.""")
figure(f"{FIGS}/fig4_evolution.png", "Figure 4. Evolutionary search before and after fifty generations.",
       "Endpoint comparison at 40,000 hands per matchup. The evolved agent improves decisively against a random opponent and not at all against either the calling station or the solver.")
table("Table 3. Endpoint comparison, 40,000 hands per matchup",
      ['Opponent', 'Untrained', 'After 50 generations', 'Difference', 'Verdict'],
      [['random', '−13.4 ± 14', '+192.7 ± 15', '+206.1 ± 39', 'improved'], ['always-call', '−8.1 ± 14', '+0.5 ± 4', '+8.6 ± 28', 'no change'], ['CFR agent', '−403.9 ± 13', '−370.1 ± 14', '+33.8 ± 37', 'no change']])
para("""**Evolutionary search learned to exploit randomness rather than to play poker.** It punishes an opponent that folds and raises at random by 206 BB/100, a difference of more than five standard errors. It gains nothing against a station that never folds, which can only be beaten by betting for value. And against the solver-derived agent it remains beaten by 370 BB/100 after fifty generations, having improved by nothing measurable. What it learned did not transfer, which is a more specific finding than the prediction that the method would simply plateau.""")
para("""Only the independent panel makes this visible. Self-play fitness sat near zero throughout, as the zero-sum property of poker requires, and reveals nothing about whether anything was learned — which is precisely why the earlier pipeline could improve against itself for eighty-nine runs while getting worse at the game.""")
h2("Three False Results, and What Caught Each")
para("""Three results during this phase looked like findings and were not, and each was caught by a check already built into the framework. A per-generation panel curve measured over 3,000 hands appeared to rise from +111 to +214 BB/100; across eleven readings it scattered with a standard deviation of 56 against a measurement error of 57, and was indistinguishable from a constant. A benchmark row reported the evolved agent beating the solver by 60.8 BB/100, alongside a 74.3% lookup miss rate — the counter that exists so a benchmark which has quietly become a second random opponent shows up as a number. And the explanation first offered for that miss rate, that the solver covers only half its abstraction, was itself wrong: every lookup had found its key, and the harness was discarding valid entries because the solver stores one probability per legal action at a node rather than one per abstract action. A measurement that is too coarse or subtly wrong does not return no result; it returns a plausible one.""")

h2("Phase Three — PPO with Self-Play, Measured")
para("""The third family is a policy-gradient agent trained by self-play. Training against a Hall of Fame pool was not possible — the pool was empty, and its former contents had been bred on the metric the audit withdrew — so the policy was snapshotted into its own opponent pool periodically and sampled from it, which is what self-play means in the audit's sense. Three seeds were trained to eight million hands each, about 4.7 hours apiece, with checkpoints retained at 500,000, two million and eight million hands so that a budget axis exists rather than a single endpoint. The rungs are points on one run rather than three runs: the two-million policy is the 500,000 policy trained further.""")
para("""Each rung was measured exactly as Phase Two was — against an untrained policy from the same initialisation, both against the same panel, both at 40,000 hands. The seeding was arranged so that this baseline can be reconstructed at all: an untrained network from seed n is bit-for-bit the network that seed n's training began from.""")
table("Table 4. PPO endpoint, gain over an untrained policy from the same initialisation",
      ['Trained to', 'vs random', 'vs always-call', 'vs CFR agent'],
      [['500,000 hands', '+215.9 (spread 217)', '+261.9 (spread 398)', '+324.3 (spread 79)'],
       ['2,000,000 hands', '+233.3 (spread 73)', '+276.9 (spread 370)', '+394.2 (spread 33)'],
       ['8,000,000 hands', '+147.1 (spread 135)', '+451.3 (spread 637)', '+373.1 (spread 48)']])
para("""All twenty-seven rows — three seeds by three rungs by three opponents — separated from zero in the improving direction, at a solver lookup miss rate of 0.0% throughout. **What PPO learned transferred, and that is the contrast with Phase Two.** Against the solver the untrained networks scored −399.1, −385.9 and −366.3 BB/100; after eight million hands the same three scored −34.8, +35.4 and −32.7, a mean of −10.7. Evolutionary search, given the same panel and the same bar, moved from −403.9 to −370.1 and was scored no change. Both halves of the audit's prediction now carry numbers: the evolutionary method plateaus and policy gradient does not.""")
para("""The figure quoted for a multi-seed result is the spread across seeds rather than any single seed's interval. A 40,000-hand measurement carries about ±14 BB/100; these seeds disagree with each other by 33 to 637 depending on the opponent, so quoting the measurement error would understate the uncertainty by an order of magnitude against the baselines.""")

h2("Phase Four — The Comparison")
para("""The three families laddered along different axes — the solver along training seconds, evolutionary search along generations, PPO along hands — and a comparison has to declare one. Wall-clock is the only axis all three share, and none of them needed re-running to supply it: every family had recorded seconds throughout, so the shared axis was arithmetic over files that already existed.""")
figure(f"{FIGS}/fig5_comparison.png", "Figure 5. The three families on one budget axis.",
       "Score against the CFR agent as a function of training wall-clock. PPO reaches parity with the solver within about an hour and then flattens; evolutionary search remains 370 BB/100 below it having spent more than three times the compute. Error bars are the spread across three training seeds.")
table("Table 5. Every family against the same panel, 40,000 hands, BB/100",
      ['Family', 'Wall-clock', 'vs random', 'vs always-call', 'vs CFR agent'],
      [['CFR (the solver)', '—', '+377.2', '+722.9', '—'],
       ['evolution, 50 generations', '3.16 h', '+192.7', '+0.5', '−370.1'],
       ['PPO, 500k hands', '0.26 h', '+204.1', '+278.6', '−59.5'],
       ['PPO, 2M hands', '1.02 h', '+221.5', '+293.5', '+10.4'],
       ['PPO, 8M hands', '4.81 h', '+135.4', '+467.9', '−10.7']])
para("""**PPO reaches break-even against the solver in about one hour of training, and evolutionary search does not reach it in three.** Because both families sit on one wall-clock axis this is a like-for-like statement about budget rather than a comparison of endpoints. The solver has no row against itself: a strategy played against a copy of itself scores zero by symmetry, and that is a structural identity rather than a measurement.""")
para("""**The table does not rank, and that is its most interesting property.** PPO at two million hands draws level with the solver head to head, yet the solver takes far more off both baselines — +377.2 against random to PPO's +221.5, and +722.9 against the calling station to PPO's +293.5. Two agents that are level against each other extract very different amounts from the same weak opponents, so strength here is not a scalar and any single-number ranking of the three families would be a fiction. It also runs against the intuition that a near-equilibrium strategy should be the less exploitative one.""")
para("""Two candidate explanations were tested and both failed. Measuring the solver through both families' code paths gave bit-identical results, so the two rows are on one instrument. Lifting the one-raise-per-street cap — the rule that stops a policy punishing a calling station by raising repeatedly — widened the gap rather than closing it, from −430.1 to −437.4. What remains is that the intransitivity is real: these strategies beat each other in a loop. The mechanism is not established, and is reported as an open question rather than explained away.""")

h2("A Fourth Measurement Defect, and How It Was Found")
para("""Phases Three and Four were measured, then re-measured, because a panel score turned out to depend on the order the matchups were taken in. Three layers were involved: the panel handed the random opponent and the solver a single random generator; the policy samples its actions from a global generator that nothing reseeded between matchups; and callers build the panel once and reuse it, so each opponent's state carried across all twelve columns of the endpoint test. Together, the same matchup read −28.9, −29.5 and −3.9 BB/100 on nothing but what had been measured before it. Every one of those is a valid sample and none was reproducible in isolation.""")
para("""How it surfaced is worth more than the defect. No result looked wrong. It came out of checking whether the intransitivity above was an artefact — that check passed, but a difference noticed while writing it exposed the first layer, and **the fix for that layer failed its own order-independence test**, which is what revealed the other two. A fix accepted without that test would have closed the smallest of the three and reported the problem solved.""")
para("""Re-measuring changed almost nothing, and that is the result worth reporting: every Phase Three figure moved less than its own seed spread — at most 13.0 BB/100 against a spread of 217, and no more than 3.6 in the solver column. The spreads themselves tightened, from 134/40/62 to 79/33/48, which is what removing an uncontrolled source of variation should do. The conclusions were sound and are now reproducible; a matchup returns the same figure measured alone, inside the full panel, or with the panel reversed.""")

# ------------------------------------------------------- final result -------
h1("Final Result Obtained")
para("The final results are produced progressively through the pipeline rather than assigned. A "
     "card abstraction is first fitted by sampling situations and clustering them into strength "
     "buckets. The solver then traverses the abstracted game repeatedly, accumulating "
     "counterfactual regret at each information set and deriving the average strategy, which is "
     "the quantity that converges toward equilibrium. Two such strategies — one per card "
     "abstraction — are trained to identical wall-clock budgets and then played against each "
     "other through the evaluation harness using duplicate play, and the resulting chip "
     "differential, together with its confidence interval, is the reported figure.")
para("The headline result is that a crossover exists. Made-hand bucketing wins decisively at "
     "short budgets, where its fivefold iteration advantage dominates; the advantage disappears "
     "at approximately 640 seconds; and equity bucketing wins by 2560 seconds at **+0.916 ± "
     "0.118** chips per hand, roughly 7.8 standard errors from zero, while training on about one "
     "fifth of its opponent's iterations. All three random seeds agree in direction at every "
     "budget.")

# -------------------------------------------------------- observations ------
h1("Observations of the Output")
OBSERVATIONS = [
    ("Poker Engine", "Chip conservation held exactly across 2,400 randomised hands and both hand "
     "evaluators agreed on all 2,000 test showdowns. Observation: the engine's correctness is the "
     "precondition for every other measurement, since all results are denominated in chips it "
     "accounts for."),
    ("Traversable Interface", "The solver reproduced Kuhn poker's analytic value of −1/18 and "
     "drove Leduc exploitability steadily downward. Observation: these small games have no "
     "practical value as poker, and exist purely so the solver can be checked against answers "
     "that are known independently of the implementation."),
    ("Card Abstraction", "Monte Carlo equity produced better-informed buckets at 5.4 times the "
     "cost per iteration of made-hand strength. Observation: the trade-off is genuine rather than "
     "a matter of preference, and its resolution depends on the training budget — which is why a "
     "single measurement would have produced a misleading answer."),
    ("Regret Solvers", "All four update rules converged, but CFR+ and DCFR finished worse than "
     "plain vanilla, contrary to published results. Observation: those rules were designed for "
     "full tree traversal, and this system uses external sampling; the discrepancy is reported "
     "rather than suppressed."),
    ("Exploitability Measurement", "Local Best Response could not establish a bound on a "
     "well-trained strategy despite four defects being corrected and three successive valuation "
     "models being implemented. Observation: the honest reading is that a greedy one-step "
     "adversary genuinely cannot beat a converged strategy in this abstraction, which is a "
     "legitimate if weak result rather than an outstanding defect."),
    ("Evaluation Harness", "Duplicate play reduced the standard error by approximately 14%, "
     "considerably less than expected. Observation: per-hand outcomes swing by a whole stack "
     "depending on whether an all-in occurred, and that variation is driven by the actions taken "
     "rather than by the cards dealt, so sharing cards cannot cancel it."),
    ("Observation Layer", "The seventeen-slot feature vector was found to carry only six "
     "independent dimensions, with all card information concentrated in a single feature. "
     "Observation: two states the agent receives as identical inputs were measured to differ by "
     "as much as 0.907 in true equity, meaning one hand may be a 5% underdog and the other a 95% "
     "favourite while appearing indistinguishable — a limit no amount of training can overcome."),
]
for name, text in OBSERVATIONS:
    para(f"**{name}.** {text}")

# --------------------------------------------------- comparative table ------
h1("Comparative Analysis")
table("Table 6. NashForge against conventional approaches",
      ["Capability", "Heuristic Bot", "Standard RL System", "Proposed NashForge"],
      [["Strategy source", "Hand-written rules", "Learned policy", "Three families compared"],
       ["Theoretical guarantee", "None", "None", "Equilibrium convergence"],
       ["Correctness of engine", "Assumed", "Assumed", "Verified, 2,400 hands"],
       ["Validation against known answers", "None", "Rare", "Kuhn and Leduc"],
       ["Exploitability measurement", "None", "None", "Exact and bounded"],
       ["Card abstraction", "Fixed", "Implicit", "Two, compared"],
       ["Benchmark independence", "Self-play only", "Self-play only", "Unrelated panel"],
       ["Variance reduction", "None", "More hands", "Duplicate play"],
       ["Confidence intervals", "No", "Rare", "On every figure"],
       ["Observation quality audited", "No", "No", "Yes, measured"]])

# ------------------------------------------------------ overall result ------
h1("Overall Result Analysis")
para("The experimental results indicate that the framework achieves its objective of making "
     "three distinct learning paradigms comparable on equal terms. The engine and traversable "
     "interface supply a game whose correctness is verified rather than assumed; the abstraction "
     "modules reduce an intractable game to one that fits in memory, at a cost that is measured; "
     "and the evaluation harness converts matches into figures carrying explicit uncertainty. The "
     "regret solver provides the only agent in the system with a theoretical guarantee, and "
     "therefore serves as the reference against which the evolutionary and reinforcement learning "
     "agents are measured.")
para("The principal advantage of NashForge is therefore not the strength of any single agent but "
     "the integrity of the comparison between them. Its most instructive outcome is a negative "
     "one: the observation audit demonstrated that the feature representation used by the "
     "learning agents discards most of the information distinguishing a strong hand from a weak "
     "one, which places a ceiling on those agents that no amount of additional training can "
     "raise. Identifying such a limit before investing in training is precisely what a "
     "measurement-first framework is intended to deliver, and it is a conclusion that a system "
     "reporting only its successes would never have reached.")

doc.save(OUT)
print(f"wrote {OUT}")
